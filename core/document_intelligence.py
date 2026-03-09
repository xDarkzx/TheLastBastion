"""
Document Intelligence — From Raw Bytes to Structured Fields.

Converts ingested raw data into normalized, schema-conformant
structured fields that the verification stack can validate.

Extraction pipeline:
    1. Format-specific extraction (PDF tables, image OCR, etc.)
    2. LLM-assisted field mapping ("which field is the price?")
    3. Schema inference from content patterns
    4. Confidence scoring for each extracted field
    5. Output -> CleanedData ready for verification

This module does NOT verify data — it only EXTRACTS it.
Verification is handled by the verification stack.
"""
import io
import json
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

try:
    import fitz as pymupdf  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import docx as python_docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger("DocumentIntelligence")


@dataclass
class ExtractedField:
    """A single field extracted from a document."""
    name: str                       # Field name (e.g., "company", "price")
    value: Any                      # Extracted value
    confidence: float               # 0.0 to 1.0 — how sure we are
    source_region: str = ""         # Where in the document (page, cell, line)
    data_type: str = "string"       # Inferred type: string, number, date, currency
    raw_text: str = ""              # Original text before normalization

    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "value": self.value,
            "confidence": round(self.confidence, 4),
            "data_type": self.data_type,
            "raw_text": self.raw_text,
        }


@dataclass
class CleanedData:
    """
    Fully extracted and normalized data from a document.

    This is the output that goes into the verification stack.
    """
    submission_id: str
    fields: Dict[str, Any]              # name -> value
    field_details: List[ExtractedField]  # Full extraction details
    inferred_schema: Dict[str, str]      # name -> expected data type
    overall_confidence: float            # Average field confidence
    document_type: str = ""              # invoice, receipt, manual, report...
    extraction_method: str = ""          # ocr, table, json_parse, llm_extract
    warnings: List[str] = field(default_factory=list)
    extracted_at: str = ""

    def __post_init__(self) -> None:
        if not self.extracted_at:
            self.extracted_at = datetime.utcnow().isoformat()

    def to_dict(self) -> Dict[str, Any]:
        return {
            "submission_id": self.submission_id,
            "fields": self.fields,
            "field_details": [f.to_dict() for f in self.field_details],
            "inferred_schema": self.inferred_schema,
            "overall_confidence": round(self.overall_confidence, 4),
            "document_type": self.document_type,
            "extraction_method": self.extraction_method,
            "warnings": self.warnings,
        }


# Common field patterns for auto-detection
FIELD_PATTERNS: Dict[str, Dict[str, Any]] = {
    "date": {
        "regex": r"\b\d{4}-\d{2}-\d{2}\b|\b\d{2}/\d{2}/\d{4}\b|\b\d{2}-\d{2}-\d{4}\b",
        "type": "date",
    },
    "currency": {
        "regex": r"\$[\d,]+\.?\d{0,2}|\b\d+\.\d{2}\s*(?:NZD|AUD|USD|GBP)\b",
        "type": "currency",
    },
    "email": {
        "regex": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
        "type": "email",
    },
    "phone": {
        "regex": r"\b(?:\+?64|0)\s*\d[\d\s-]{7,12}\b",
        "type": "phone",
    },
    "percentage": {
        "regex": r"\b\d+\.?\d*\s*%",
        "type": "percentage",
    },
    "gst_number": {
        "regex": r"\b\d{2,3}-?\d{3}-?\d{3}\b",
        "type": "gst_number",
    },
}


class DocumentIntelligence:
    """
    Extracts structured fields from raw documents.

    Supports multiple extraction strategies:
    - JSON: Direct parse (highest confidence)
    - CSV: Row/column mapping
    - Text: Pattern matching + heuristics
    - PDF: Text extraction via PyMuPDF, fallback to OCR
    - Image: OCR via Tesseract (pytesseract)

    The LLM is used for field mapping ONLY — not for
    generating or fabricating data.
    """

    def __init__(self) -> None:
        self._extraction_count = 0

    def extract_from_json(
        self, data: Dict[str, Any], submission_id: str
    ) -> CleanedData:
        """
        Extracts fields from pre-structured JSON.

        This is the highest confidence path — the data is
        already structured, we just need to infer types.
        """
        field_details = []
        inferred_schema = {}

        for key, value in data.items():
            dtype = self._infer_type(key, value)
            inferred_schema[key] = dtype

            field_details.append(ExtractedField(
                name=key,
                value=value,
                confidence=1.0,    # JSON = already structured
                data_type=dtype,
                raw_text=str(value),
            ))

        doc_type = self._infer_document_type(data)
        self._extraction_count += 1

        logger.info(
            f"EXTRACT (JSON): {submission_id} — "
            f"{len(field_details)} fields, type={doc_type}"
        )

        return CleanedData(
            submission_id=submission_id,
            fields=data,
            field_details=field_details,
            inferred_schema=inferred_schema,
            overall_confidence=1.0,
            document_type=doc_type,
            extraction_method="json_parse",
        )

    def extract_from_text(
        self, text: str, submission_id: str
    ) -> CleanedData:
        """
        Extracts fields from unstructured text using pattern matching.

        Uses regex patterns to find dates, currencies, emails,
        phone numbers, percentages, and GST numbers.
        """
        field_details = []
        fields = {}
        inferred_schema = {}

        for pattern_name, pattern_def in FIELD_PATTERNS.items():
            matches = re.findall(pattern_def["regex"], text)
            if matches:
                for i, match in enumerate(matches):
                    field_name = (
                        pattern_name if len(matches) == 1
                        else f"{pattern_name}_{i+1}"
                    )
                    normalized = self._normalize_value(
                        match, pattern_def["type"]
                    )

                    field_details.append(ExtractedField(
                        name=field_name,
                        value=normalized,
                        confidence=0.75,    # Pattern match confidence
                        data_type=pattern_def["type"],
                        raw_text=match,
                    ))
                    fields[field_name] = normalized
                    inferred_schema[field_name] = pattern_def["type"]

        # Also try to extract key:value pairs
        kv_matches = re.findall(
            r"(?:^|\n)\s*([A-Za-z][\w\s]{1,30}):\s*(.+?)(?:\n|$)",
            text,
        )
        for key, value in kv_matches:
            clean_key = key.strip().lower().replace(" ", "_")
            if clean_key not in fields:
                dtype = self._infer_type(clean_key, value.strip())
                normalized = self._normalize_value(value.strip(), dtype)

                field_details.append(ExtractedField(
                    name=clean_key,
                    value=normalized,
                    confidence=0.6,     # KV pair = moderate
                    data_type=dtype,
                    raw_text=value.strip(),
                ))
                fields[clean_key] = normalized
                inferred_schema[clean_key] = dtype

        overall_conf = (
            sum(f.confidence for f in field_details) / len(field_details)
            if field_details else 0.0
        )

        self._extraction_count += 1

        logger.info(
            f"EXTRACT (text): {submission_id} — "
            f"{len(field_details)} fields extracted"
        )

        return CleanedData(
            submission_id=submission_id,
            fields=fields,
            field_details=field_details,
            inferred_schema=inferred_schema,
            overall_confidence=round(overall_conf, 4),
            document_type="text_document",
            extraction_method="pattern_match",
        )

    def extract_from_csv_rows(
        self,
        rows: List[Dict[str, str]],
        submission_id: str,
    ) -> CleanedData:
        """
        Extracts and normalizes fields from CSV rows.

        Infers types per column and normalizes values.
        """
        if not rows:
            return CleanedData(
                submission_id=submission_id,
                fields={},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                extraction_method="csv_parse",
                warnings=["No rows to extract"],
            )

        # Infer schema from all rows
        inferred_schema = {}
        for col in rows[0].keys():
            sample_values = [row.get(col, "") for row in rows[:10]]
            inferred_schema[col] = self._infer_column_type(sample_values)

        # Normalize all values
        normalized_rows = []
        field_details = []
        for row_idx, row in enumerate(rows):
            normalized_row = {}
            for col, value in row.items():
                dtype = inferred_schema.get(col, "string")
                normalized = self._normalize_value(value, dtype)
                normalized_row[col] = normalized

                field_details.append(ExtractedField(
                    name=col,
                    value=normalized,
                    confidence=0.85,    # CSV = well-structured
                    source_region=f"row_{row_idx + 1}",
                    data_type=dtype,
                    raw_text=str(value),
                ))
            normalized_rows.append(normalized_row)

        self._extraction_count += 1

        logger.info(
            f"EXTRACT (CSV): {submission_id} — "
            f"{len(rows)} rows × {len(inferred_schema)} columns"
        )

        return CleanedData(
            submission_id=submission_id,
            fields={"rows": normalized_rows, "row_count": len(rows)},
            field_details=field_details,
            inferred_schema=inferred_schema,
            overall_confidence=0.85,
            document_type="tabular_data",
            extraction_method="csv_parse",
        )

    def extract_from_image(
        self, image_bytes: bytes, submission_id: str
    ) -> CleanedData:
        """
        Extracts text from an image using Tesseract OCR,
        then runs pattern matching on the extracted text.
        """
        if not HAS_PIL:
            logger.warning("EXTRACT (image): PIL not installed, returning empty")
            return CleanedData(
                submission_id=submission_id,
                fields={},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                extraction_method="ocr_unavailable",
                warnings=["PIL/Pillow not installed"],
            )

        if not HAS_TESSERACT:
            logger.warning("EXTRACT (image): pytesseract not installed, returning empty")
            return CleanedData(
                submission_id=submission_id,
                fields={},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                extraction_method="ocr_unavailable",
                warnings=["pytesseract not installed"],
            )

        try:
            image = Image.open(io.BytesIO(image_bytes))
            raw_text = pytesseract.image_to_string(image)
        except Exception as e:
            logger.error(f"EXTRACT (image): OCR failed: {e}")
            return CleanedData(
                submission_id=submission_id,
                fields={},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                extraction_method="ocr_failed",
                warnings=[f"OCR error: {e}"],
            )

        if not raw_text.strip():
            logger.info(f"EXTRACT (image): {submission_id} — no text found")
            return CleanedData(
                submission_id=submission_id,
                fields={"raw_text": ""},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                document_type="image_empty",
                extraction_method="ocr",
                warnings=["No text detected in image"],
            )

        # OCR succeeded — run text extraction on the result
        result = self.extract_from_text(raw_text, submission_id)
        result.extraction_method = "ocr"
        # Lower confidence because OCR introduces error
        result.overall_confidence = round(result.overall_confidence * 0.85, 4)
        for fd in result.field_details:
            fd.confidence = round(fd.confidence * 0.85, 4)

        # Store raw OCR text for audit
        result.fields["_ocr_raw_text"] = raw_text

        self._extraction_count += 1
        logger.info(
            f"EXTRACT (OCR): {submission_id} — "
            f"{len(result.field_details)} fields from {len(raw_text)} chars"
        )
        return result

    def extract_from_pdf(
        self, pdf_bytes: bytes, submission_id: str
    ) -> CleanedData:
        """
        Extracts text from a PDF using PyMuPDF (text-layer first),
        falls back to OCR on each page if no text found.
        """
        if not HAS_PYMUPDF:
            # Fallback: try OCR on the raw bytes (treat as image)
            if HAS_PIL and HAS_TESSERACT:
                logger.info("EXTRACT (PDF): PyMuPDF not available, trying OCR fallback")
                return self.extract_from_image(pdf_bytes, submission_id)
            logger.warning("EXTRACT (PDF): No PDF/OCR libs available")
            return CleanedData(
                submission_id=submission_id,
                fields={},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                extraction_method="pdf_unavailable",
                warnings=["PyMuPDF not installed"],
            )

        try:
            doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
        except Exception as e:
            logger.error(f"EXTRACT (PDF): Open failed: {e}")
            return CleanedData(
                submission_id=submission_id,
                fields={},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                extraction_method="pdf_failed",
                warnings=[f"PDF open error: {e}"],
            )

        page_texts = []
        ocr_pages = 0
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")

            if text.strip():
                page_texts.append(text)
            elif HAS_PIL and HAS_TESSERACT:
                # No text layer — OCR the page image
                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(img)
                if ocr_text.strip():
                    page_texts.append(ocr_text)
                    ocr_pages += 1

        doc.close()

        combined_text = "\n\n".join(page_texts)
        if not combined_text.strip():
            return CleanedData(
                submission_id=submission_id,
                fields={"raw_text": ""},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                document_type="pdf_empty",
                extraction_method="pdf_extract",
                warnings=["No text extracted from PDF"],
            )

        result = self.extract_from_text(combined_text, submission_id)
        result.extraction_method = "pdf_extract" if ocr_pages == 0 else "pdf_ocr_hybrid"

        # OCR pages reduce confidence
        if ocr_pages > 0:
            ocr_ratio = ocr_pages / max(len(page_texts), 1)
            penalty = 1.0 - (ocr_ratio * 0.15)
            result.overall_confidence = round(result.overall_confidence * penalty, 4)

        result.fields["_pdf_pages"] = len(page_texts)
        result.fields["_ocr_pages"] = ocr_pages
        result.fields["_raw_text"] = combined_text

        self._extraction_count += 1
        logger.info(
            f"EXTRACT (PDF): {submission_id} — "
            f"{len(page_texts)} pages ({ocr_pages} via OCR), "
            f"{len(result.field_details)} fields"
        )
        return result

    def extract_from_excel(
        self, excel_bytes: bytes, submission_id: str
    ) -> CleanedData:
        """
        Extracts structured data from Excel (.xlsx) files using openpyxl.

        Reads all sheets, converts each to rows with headers from row 1,
        then processes as structured tabular data.
        """
        if not HAS_OPENPYXL:
            logger.warning("EXTRACT (Excel): openpyxl not installed")
            return CleanedData(
                submission_id=submission_id,
                fields={},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                extraction_method="excel_unavailable",
                warnings=["openpyxl not installed"],
            )

        try:
            wb = openpyxl.load_workbook(io.BytesIO(excel_bytes), read_only=True, data_only=True)
        except Exception as e:
            logger.error(f"EXTRACT (Excel): Open failed: {e}")
            return CleanedData(
                submission_id=submission_id,
                fields={},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                extraction_method="excel_failed",
                warnings=[f"Excel open error: {e}"],
            )

        all_rows = []
        field_details = []
        inferred_schema = {}
        sheet_count = 0

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_iter = ws.iter_rows(values_only=True)

            # First row = headers
            try:
                header_row = next(rows_iter)
            except StopIteration:
                continue

            headers = [str(h).strip() if h is not None else f"col_{i}" for i, h in enumerate(header_row)]
            sheet_count += 1

            for row in rows_iter:
                if all(cell is None for cell in row):
                    continue  # Skip empty rows
                row_dict = {}
                for i, cell_value in enumerate(row):
                    if i < len(headers):
                        col_name = headers[i]
                        value = cell_value if cell_value is not None else ""
                        row_dict[col_name] = value
                if row_dict:
                    all_rows.append(row_dict)

        wb.close()

        if not all_rows:
            return CleanedData(
                submission_id=submission_id,
                fields={"_sheet_count": sheet_count},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                document_type="excel_empty",
                extraction_method="excel_extract",
                warnings=["No data rows found in Excel file"],
            )

        # Infer schema from first row's headers + sample values
        for col in all_rows[0].keys():
            sample_values = [str(row.get(col, "")) for row in all_rows[:10]]
            inferred_schema[col] = self._infer_column_type(sample_values)

        # Build field details for each column (summary across all rows)
        for col in all_rows[0].keys():
            sample = all_rows[0].get(col, "")
            dtype = inferred_schema.get(col, "string")
            field_details.append(ExtractedField(
                name=col,
                value=sample,
                confidence=0.90,
                data_type=dtype,
                raw_text=str(sample),
                source_region=f"sheet={sheet_count} rows={len(all_rows)}",
            ))

        # Flatten to fields dict
        fields = {
            "_rows": all_rows,
            "_row_count": len(all_rows),
            "_sheet_count": sheet_count,
            "_columns": list(all_rows[0].keys()) if all_rows else [],
        }
        # Add first-row values as top-level fields
        if all_rows:
            for key, value in all_rows[0].items():
                fields[key] = value

        overall_conf = 0.90  # Excel data is already structured
        doc_type = self._infer_document_type(fields)
        self._extraction_count += 1

        logger.info(
            f"EXTRACT (Excel): {submission_id} — "
            f"{sheet_count} sheet(s), {len(all_rows)} rows, "
            f"{len(field_details)} columns"
        )

        return CleanedData(
            submission_id=submission_id,
            fields=fields,
            field_details=field_details,
            inferred_schema=inferred_schema,
            overall_confidence=overall_conf,
            document_type=doc_type,
            extraction_method="excel_extract",
        )

    def extract_from_docx(
        self, docx_bytes: bytes, submission_id: str
    ) -> CleanedData:
        """
        Extracts text from Word (.docx) files using python-docx.

        Reads all paragraphs and tables, then runs pattern matching
        on the combined text.
        """
        if not HAS_DOCX:
            logger.warning("EXTRACT (Word): python-docx not installed")
            return CleanedData(
                submission_id=submission_id,
                fields={},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                extraction_method="docx_unavailable",
                warnings=["python-docx not installed"],
            )

        try:
            doc = python_docx.Document(io.BytesIO(docx_bytes))
        except Exception as e:
            logger.error(f"EXTRACT (Word): Open failed: {e}")
            return CleanedData(
                submission_id=submission_id,
                fields={},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                extraction_method="docx_failed",
                warnings=[f"Word document open error: {e}"],
            )

        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract tables
        table_data = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):
                    table_data.append(row_cells)

        combined_text = "\n".join(paragraphs)
        if table_data:
            table_text = "\n".join([" | ".join(row) for row in table_data])
            combined_text += "\n\n" + table_text

        if not combined_text.strip():
            return CleanedData(
                submission_id=submission_id,
                fields={"raw_text": ""},
                field_details=[],
                inferred_schema={},
                overall_confidence=0.0,
                document_type="docx_empty",
                extraction_method="docx_extract",
                warnings=["No text found in Word document"],
            )

        # Run text extraction pipeline on the combined text
        result = self.extract_from_text(combined_text, submission_id)
        result.extraction_method = "docx_extract"
        # Word docs have good text fidelity — slight confidence boost vs raw text
        result.overall_confidence = round(min(result.overall_confidence * 1.1, 0.95), 4)

        # Add document metadata
        result.fields["_paragraph_count"] = len(paragraphs)
        result.fields["_table_count"] = len(doc.tables)
        result.fields["_raw_text"] = combined_text

        self._extraction_count += 1
        logger.info(
            f"EXTRACT (Word): {submission_id} — "
            f"{len(paragraphs)} paragraphs, {len(doc.tables)} tables, "
            f"{len(result.field_details)} fields"
        )
        return result

    def extract_auto(
        self, raw_bytes: bytes, file_type: str, submission_id: str
    ) -> CleanedData:
        """
        Auto-routes to the right extraction method based on file type.
        """
        ft = file_type.lower()
        if ft in ("pdf", "application/pdf"):
            return self.extract_from_pdf(raw_bytes, submission_id)
        elif ft in ("image/jpeg", "image/png", "jpeg", "jpg", "png", "image",
                     "image/bmp", "image/tiff", "image/webp", "bmp", "tiff", "webp"):
            return self.extract_from_image(raw_bytes, submission_id)
        elif ft in ("excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     "xlsx", "xls"):
            return self.extract_from_excel(raw_bytes, submission_id)
        elif ft in ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     "word", "doc"):
            return self.extract_from_docx(raw_bytes, submission_id)
        elif ft in ("csv", "text/csv"):
            text = raw_bytes.decode("utf-8", errors="replace")
            import csv as csv_mod
            reader = csv_mod.DictReader(io.StringIO(text))
            rows = list(reader)
            return self.extract_from_csv_rows(rows, submission_id)
        elif ft in ("json", "application/json"):
            data = json.loads(raw_bytes.decode("utf-8", errors="replace"))
            if isinstance(data, dict):
                return self.extract_from_json(data, submission_id)
        # Default: treat as plain text
        text = raw_bytes.decode("utf-8", errors="replace")
        return self.extract_from_text(text, submission_id)

    def _infer_type(self, key: str, value: Any) -> str:
        """Infers the data type from key name and value."""
        if isinstance(value, bool):
            return "boolean"
        if isinstance(value, int):
            return "integer"
        if isinstance(value, float):
            return "number"
        if isinstance(value, list):
            return "array"
        if isinstance(value, dict):
            return "object"

        # String value — check key name for hints
        key_lower = key.lower()
        date_keys = {"date", "created", "updated", "timestamp", "time", "day"}
        price_keys = {"price", "cost", "total", "amount", "rate", "fee"}
        id_keys = {"id", "number", "code", "gst", "nzbn", "ref"}

        if any(dk in key_lower for dk in date_keys):
            return "date"
        if any(pk in key_lower for pk in price_keys):
            return "currency"
        if any(ik in key_lower for ik in id_keys):
            return "identifier"

        return "string"

    def _infer_column_type(self, values: List[str]) -> str:
        """Infers column type from sample values."""
        non_empty = [v for v in values if v.strip()]
        if not non_empty:
            return "string"

        # Check if all values are numeric
        numeric_count = sum(
            1 for v in non_empty
            if re.match(r"^-?\d+\.?\d*$", v.strip())
        )
        if numeric_count == len(non_empty):
            if all("." in v for v in non_empty):
                return "number"
            return "integer"

        # Check if all values are dates
        date_count = sum(
            1 for v in non_empty
            if re.match(r"\d{4}-\d{2}-\d{2}", v.strip())
        )
        if date_count == len(non_empty):
            return "date"

        return "string"

    def _normalize_value(self, value: str, dtype: str) -> Any:
        """Normalizes a raw string value to its inferred type."""
        value = value.strip()
        if dtype in ("number", "currency"):
            cleaned = re.sub(r"[^\d.-]", "", value)
            try:
                return float(cleaned) if "." in cleaned else int(cleaned)
            except ValueError:
                return value
        if dtype == "integer":
            cleaned = re.sub(r"[^\d-]", "", value)
            try:
                return int(cleaned)
            except ValueError:
                return value
        if dtype == "percentage":
            cleaned = re.sub(r"[^\d.]", "", value)
            try:
                return float(cleaned)
            except ValueError:
                return value
        return value

    def _infer_document_type(self, data: Dict[str, Any]) -> str:
        """Infers document type from field names using substring matching."""
        keys_lower = {k.lower() for k in data.keys()}
        keys_joined = " ".join(keys_lower)

        # Check for energy-specific fields first (more specific)
        energy_markers = ["electricity_price", "gas_price", "energy", "kwh"]
        if any(m in keys_joined for m in energy_markers):
            return "energy_pricing"

        invoice_markers = ["invoice_number", "invoice", "bill", "due_date"]
        if any(m in keys_joined for m in invoice_markers):
            return "invoice"

        if keys_lower & {"receipt", "transaction"}:
            return "receipt"
        if keys_lower & {"premium", "coverage", "policy"}:
            return "insurance"
        if keys_lower & {"part_number", "torque", "specification"}:
            return "technical_manual"
        if keys_lower & {"company", "provider", "price"}:
            return "pricing_data"

        return "general"
